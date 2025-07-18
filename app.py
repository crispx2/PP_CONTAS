from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from io import BytesIO
import json
import os
import sqlite3

app = Flask(__name__)

# Banco de dados
DB_PATH = 'cambistas.db'

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS cambistas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        codigo TEXT NOT NULL,
        nome TEXT NOT NULL,
        area TEXT NOT NULL
    )''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS registros_comissao (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        codigo TEXT NOT NULL,
        comissao REAL NOT NULL,
        complemento REAL NOT NULL,
        area TEXT NOT NULL
    )''')
    conn.commit()
    conn.close()

init_db()

# Caminho do arquivo de metas
META_FILE = 'metas.json'

def carregar_meta_data():
    if os.path.exists(META_FILE):
        with open(META_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def salvar_meta_data():
    with open(META_FILE, 'w', encoding='utf-8') as f:
        json.dump(meta_data_store, f, ensure_ascii=False, indent=2)

meta_data_store = carregar_meta_data()

@app.route('/')
def selecao():
    return render_template('selecao.html')

@app.route('/principal')
def principal():
    modo = request.args.get('modo', 'comissao')
    if modo == 'meta':
        return render_template('meta.html')
    else:
        return render_template('index.html', modo=modo)

# --- API Comissão ---
@app.route('/adicionar', methods=['POST'])
def adicionar():
    data = request.json
    codigo = data['codigo']
    comissao = float(data['comissao'])
    complemento = float(data['Complemento_De_Dezena'])
    area = data['area']

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO registros_comissao (codigo, comissao, complemento, area) VALUES (?, ?, ?, ?)",
                   (codigo, comissao, complemento, area))
    conn.commit()
    conn.close()

    return jsonify({'status': 'ok'})

@app.route('/dados')
def dados():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''SELECT r.codigo, c.nome, r.comissao, r.complemento, r.area
                      FROM registros_comissao r
                      JOIN cambistas c ON r.codigo = c.codigo''')
    rows = cursor.fetchall()
    conn.close()

    result = [
        {'codigo': r[0], 'cambista': r[1], 'comissao': r[2], 'Complemento_De_Dezena': r[3], 'area': r[4]}
        for r in rows
    ]
    return jsonify(result)

@app.route('/apagar/<int:index>', methods=['DELETE'])
def apagar(index):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM registros_comissao WHERE id = ?", (index,))
    conn.commit()
    conn.close()
    return jsonify({'status': 'ok'})

@app.route('/cadastrar-cambista', methods=['POST'])
def cadastrar_cambista():
    data = request.json
    codigo = data['codigo']
    nome = data['nome']
    area = data['area']

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("INSERT INTO cambistas (codigo, nome, area) VALUES (?, ?, ?)", (codigo, nome, area))
    conn.commit()
    conn.close()

    return jsonify({'status': 'ok'})

@app.route('/cambistas')
def listar_cambistas():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT codigo, nome, area FROM cambistas")
    rows = cursor.fetchall()
    conn.close()
    return jsonify([{'codigo': r[0], 'nome': r[1], 'area': r[2]} for r in rows])

@app.route('/exportar-word', methods=['POST'])
def exportar_word():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''SELECT r.codigo, c.nome, r.comissao, r.complemento, r.area
                      FROM registros_comissao r
                      JOIN cambistas c ON r.codigo = c.codigo
                      ORDER BY r.area''')
    rows = cursor.fetchall()
    conn.close()

    doc = Document()
    doc.add_heading('Relatório de Cambistas por Área', 0)

    areas = sorted(set(r[4] for r in rows))
    for area in areas:
        doc.add_heading(area, level=1)
        tabela_area = [r for r in rows if r[4] == area]

        if not tabela_area:
            doc.add_paragraph('Nenhum registro.')
            continue

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Código'
        hdr_cells[1].text = 'Cambista'
        hdr_cells[2].text = 'Comissão (R$)'
        hdr_cells[3].text = 'Complemento de Dezena (R$)'

        for item in tabela_area:
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = f"R$ {item[2]:,.2f}"
            row_cells[3].text = f"R$ {item[3]:,.2f}"

        total_comissao = sum(x[2] for x in tabela_area)
        total_complemento = sum(x[3] for x in tabela_area)
        doc.add_paragraph(f"Total Comissão: R$ {total_comissao:,.2f}")
        doc.add_paragraph(f"Total Complemento: R$ {total_complemento:,.2f}")
        doc.add_paragraph('')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='Relatorio_Cambistas.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# --- Metas ---
@app.route('/adicionar-meta', methods=['POST'])
def adicionar_meta():
    data = request.json
    meta_data_store.append(data)
    salvar_meta_data()
    return jsonify({'status': 'ok'})

@app.route('/dados-meta')
def dados_meta():
    return jsonify(meta_data_store)

@app.route('/apagar-meta/<int:index>', methods=['DELETE'])
def apagar_meta(index):
    try:
        meta_data_store.pop(index)
        salvar_meta_data()
        return jsonify({'status': 'ok'})
    except IndexError:
        return jsonify({'status': 'error', 'message': 'Índice inválido'}), 400

@app.route('/exportar-word-meta', methods=['POST'])
def exportar_word_meta():
    doc = Document()
    doc.add_heading('Relatório de Metas por Área', 0)

    areas = sorted(set(d['area'] for d in meta_data_store))
    for area in areas:
        doc.add_heading(area, level=1)
        tabela_area = [d for d in meta_data_store if d['area'] == area]

        if not tabela_area:
            doc.add_paragraph('Nenhum registro.')
            continue

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Código do Cambista'
        hdr_cells[1].text = 'Nome do Cambista'
        hdr_cells[2].text = 'Apurado (R$)'
        hdr_cells[3].text = 'Valor a Receber (R$)'

        for item in tabela_area:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('codigo', ''))
            row_cells[1].text = str(item.get('nome', ''))
            row_cells[2].text = f"R$ {float(item.get('apurado', 0)):,.2f}"
            row_cells[3].text = f"R$ {float(item.get('valorReceber', 0)):,.2f}"

        doc.add_paragraph('')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='Relatorio_Meta_Cambistas.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)